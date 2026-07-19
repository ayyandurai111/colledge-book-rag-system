"""
conftest.py
-----------
Shared fixtures for the whole test suite. Defines a deterministic
FakeSentenceTransformer that replaces sentence-transformers so tests run
fully offline, and inline sample PDF/DOCX generators (previously a
separate scripts/ module — now inlined here so tests have no external
script dependency).
"""
import sys
from pathlib import Path

import docx
import numpy as np
import pytest
from fpdf import FPDF
from fpdf.enums import XPos, YPos

sys.path.insert(0, str(Path(__file__).parent.parent / "src"))

from college_rag.embeddings.embedder import Embedder  # noqa: E402

# --------------------------------------------------------------------------- #
# Fake / deterministic embedding model — lets us test chunking/vectorstore/
# retrieval logic without downloading sentence-transformers weights.
# --------------------------------------------------------------------------- #

_TOPIC_ANCHORS = {
    "mechanics": ["newton", "force", "mass", "acceleration", "inertia", "motion"],
    "thermo": ["entropy", "heat", "thermodynamics", "disorder", "temperature"],
    "biology": ["cell", "organism", "life", "division"],
    "genetics": ["dna", "gene", "genetic", "protein", "chromosome"],
}


class FakeSentenceTransformer:
    """A fully offline, deterministic stand-in for
    sentence-transformers.SentenceTransformer.

    Produces a fixed-dimension vector for each text based on its overlap
    with a set of predefined "topic" keywords. Sentences about the same
    topic end up close together (high cosine similarity) — simulating the
    behavior of a real embedding model well enough to test chunking and
    retrieval logic.
    """

    def encode(self, texts, normalize_embeddings=True, show_progress_bar=False):
        vectors = []
        for text in texts:
            lowered = text.lower()
            vec = np.array(
                [sum(kw in lowered for kw in kws) for kws in _TOPIC_ANCHORS.values()]
                + [len(text) * 0.001],  # small tie-breaker signal
                dtype="float32",
            )
            norm = np.linalg.norm(vec)
            if norm > 0 and normalize_embeddings:
                vec = vec / norm
            vectors.append(vec)
        return np.array(vectors, dtype="float32")


@pytest.fixture
def fake_embedder() -> Embedder:
    """An Embedder instance that works without downloading a real model."""
    return Embedder(model_name="fake-test-model", model=FakeSentenceTransformer())


# --------------------------------------------------------------------------- #
# Sample document generators (inlined — no scripts/ dependency)
# --------------------------------------------------------------------------- #

def make_sample_docx(path: str) -> str:
    d = docx.Document()
    d.add_heading("Chapter 1: Mechanics", level=1)
    d.add_paragraph(
        "Newton's First Law states that an object at rest stays at rest, and an "
        "object in motion stays in motion, unless acted upon by an external force. "
        "This principle is also known as the law of inertia."
    )
    d.add_paragraph(
        "Newton's Second Law states that force equals mass times acceleration. "
        "This relationship is fundamental to classical mechanics."
    )
    d.add_heading("Chapter 2: Thermodynamics", level=1)
    d.add_paragraph(
        "The Second Law of Thermodynamics states that the total entropy of an "
        "isolated system can never decrease over time. Entropy is a measure of "
        "disorder in a system."
    )
    d.add_paragraph(
        "Heat naturally flows from hotter objects to colder objects, never the "
        "reverse, without external work being applied."
    )
    d.save(path)
    return path


def make_sample_pdf(path: str) -> str:
    pdf = FPDF()
    pdf.set_font("Helvetica", size=12)

    pdf.add_page()
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 10, "CHAPTER 1 BIOLOGY BASICS", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("Helvetica", size=12)
    pdf.multi_cell(
        0, 8,
        "A cell is the basic structural and functional unit of all living "
        "organisms. Cells are often called the building blocks of life. "
        "All cells arise from pre-existing cells through the process of cell "
        "division."
    )

    pdf.add_page()
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 10, "CHAPTER 2 GENETICS OVERVIEW", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("Helvetica", size=12)
    pdf.multi_cell(
        0, 8,
        "DNA carries the genetic instructions used in the growth, development, "
        "and reproduction of all known organisms. Genes are segments of DNA "
        "that code for specific proteins."
    )

    pdf.output(path)
    return path


@pytest.fixture
def sample_docx_path(tmp_path) -> str:
    return make_sample_docx(str(tmp_path / "sample_physics.docx"))


@pytest.fixture
def sample_pdf_path(tmp_path) -> str:
    return make_sample_pdf(str(tmp_path / "sample_biology.pdf"))
